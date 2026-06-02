import re
import tempfile
import os
import logging
from io import BytesIO
from typing import Tuple, Dict, Any, Optional, List
from datetime import datetime

logger = logging.getLogger(__name__)


MAX_CITATION_CONTENT_LENGTH = 800

MAX_FOOTNOTE_CONTENT_LENGTH = 80

FOOTNOTE_HEAD_LENGTH = 60
FOOTNOTE_TAIL_LENGTH = 20

CHINESE_NUMERALS = [
    '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
    '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十'
]


def to_chinese_numeral(n: int) -> str:
    if 1 <= n <= len(CHINESE_NUMERALS):
        return CHINESE_NUMERALS[n - 1]
    return str(n)


class ExportService:

    def export_session_to_docx(
        self,
        session_title: str,
        all_messages: List[Dict[str, Any]],
        user_message_ids: Optional[List[str]] = None,
        include_citations: bool = True,
        citation_style: str = "footnote",
    ) -> Tuple[BytesIO, str]:
        citations_map: Dict[str, Any] = {}
        for m in all_messages:
            if m.get("role") == "tool" and m.get("citations"):
                citations = m["citations"]
                if isinstance(citations, dict):
                    citations_map.update(citations)

        selected_messages = self._select_message_pairs(
            all_messages, user_message_ids
        )

        if not selected_messages:
            raise ValueError("没有可导出的对话内容")

        title = session_title or "对话导出"

        if citation_style == "footnote" and include_citations and citations_map:
            markdown_content = self._generate_chat_markdown_with_footnotes(
                title, selected_messages, citations_map, include_citations
            )
            docx_stream = self._convert_markdown_to_docx(markdown_content)
        else:
            markdown_content = self._generate_chat_markdown(
                title, selected_messages, citations_map, include_citations
            )
            docx_stream = self._convert_markdown_to_docx(markdown_content)

        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.docx"

        return docx_stream, filename

    def _select_message_pairs(
        self,
        all_messages: List[Dict[str, Any]],
        user_message_ids: Optional[List[str]] = None,
    ) -> List[Dict[str, Any]]:
        selected = []
        user_ids_set = set(user_message_ids) if user_message_ids is not None else None

        for i, m in enumerate(all_messages):
            if m['role'] != 'user':
                continue
            if user_ids_set is not None and m['id'] not in user_ids_set:
                continue

            selected.append(m)

            next_user_index = len(all_messages)
            for j in range(i + 1, len(all_messages)):
                if all_messages[j]['role'] == 'user':
                    next_user_index = j
                    break

            last_valid_assistant = None
            for j in range(i + 1, next_user_index):
                next_m = all_messages[j]
                if next_m['role'] == 'assistant' and next_m.get('content'):
                    last_valid_assistant = next_m

            if last_valid_assistant:
                selected.append(last_valid_assistant)

        return selected

    def _generate_chat_markdown_with_footnotes(
        self,
        title: str,
        messages: List[Dict],
        citations_map: Dict[str, Any],
        include_citations: bool,
    ) -> str:
        self._footnotes = []
        self._fn_id_counter = 0
        self._display_num_to_first_fn_id = {}
        self._display_num_to_first_fn_index = {}
        self._fn_index_to_display_num = {}
        self._merge_list = []
        self._original_to_export_num = {}
        self._export_num_counter = 0

        md_parts = []
        md_parts.append(f"# {title}\n")
        md_parts.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        md_parts.append(f"\n---\n")

        for msg in messages:
            role = msg['role']
            content = msg.get('content', '')

            if role == 'user':
                md_parts.append(f"\n### **用户：**\n")
                md_parts.append(f"{self._indent_paragraph(content)}\n")

            elif role == 'assistant':
                md_parts.append(f"\n### **AI助手：**\n")

                processed = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)

                if include_citations and citations_map:
                    processed = self._process_chat_citations_for_footnotes(
                        processed, citations_map
                    )
                else:
                    processed = re.sub(r'\[citation_\d+\]', '', processed)

                formatted = self._format_content_for_export(processed)
                md_parts.append(f"{formatted}\n")
                md_parts.append(f"\n\n---\n\n---\n\n")

        if self._footnotes:
            md_parts.append("\n")
            for fn_id, fn_content in self._footnotes:
                fn_content = fn_content.replace('\n', ' ').replace('\r', ' ')
                fn_content = re.sub(r'\s+', ' ', fn_content).strip()
                md_parts.append(f"[^{fn_id}]: {fn_content}\n")

        return '\n'.join(md_parts)

    def _process_chat_citations_for_footnotes(
        self,
        text: str,
        citations_map: Dict[str, Any],
    ) -> str:
        def replace_citation(match):
            citation_key = match.group(0)[1:-1]
            citation = citations_map.get(citation_key)
            if not citation:
                return ''

            original_display_num = citation.get('display_num')
            if original_display_num is None:
                return ''

            if original_display_num not in self._original_to_export_num:
                self._export_num_counter += 1
                self._original_to_export_num[original_display_num] = self._export_num_counter
            export_num = self._original_to_export_num[original_display_num]

            self._fn_id_counter += 1
            temp_fn_id = f"fn_{self._fn_id_counter}"
            current_fn_index = len(self._footnotes)
            self._fn_index_to_display_num[current_fn_index] = export_num

            if export_num in self._display_num_to_first_fn_id:
                first_fn_id = self._display_num_to_first_fn_id[export_num]
                first_fn_index = self._display_num_to_first_fn_index[export_num]
                self._merge_list.append((current_fn_index, first_fn_index, export_num))
                merge_marker = f"（同引用{export_num}）"
                self._footnotes.append((temp_fn_id, merge_marker))
            else:
                self._display_num_to_first_fn_id[export_num] = temp_fn_id
                self._display_num_to_first_fn_index[export_num] = current_fn_index
                fn_content = self._build_footnote_content(citation)
                self._footnotes.append((temp_fn_id, fn_content))

            return f"[^{temp_fn_id}]"

        result = re.sub(r'\[citation_\d+\]', replace_citation, text)
        result = re.sub(r'\]\[\^', '] [^', result)
        return result


    def _generate_chat_markdown(
        self,
        title: str,
        messages: List[Dict],
        citations_map: Dict[str, Any],
        include_citations: bool,
    ) -> str:
        self._footnotes = []
        self._fn_id_counter = 0
        self._display_num_to_first_fn_id = {}
        self._display_num_to_first_fn_index = {}
        self._fn_index_to_display_num = {}
        self._merge_list = []
        self._original_to_export_num = {}
        self._export_num_counter = 0

        md_parts = []
        md_parts.append(f"# {title}\n")
        md_parts.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        md_parts.append(f"\n---\n")

        for msg in messages:
            role = msg['role']
            content = msg.get('content', '')

            if role == 'user':
                md_parts.append(f"\n### **用户：**\n")
                md_parts.append(f"{self._indent_paragraph(content)}\n")

            elif role == 'assistant':
                md_parts.append(f"\n### **AI助手：**\n")

                processed = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL)

                if include_citations and citations_map:
                    processed = self._replace_citations_inline(processed, citations_map)
                else:
                    processed = re.sub(r'\[citation_\d+\]', '', processed)

                formatted = self._format_content_for_export(processed)
                md_parts.append(f"{formatted}\n")
                md_parts.append(f"\n\n---\n\n---\n\n")

        return '\n'.join(md_parts)

    def _replace_citations_inline(
        self, text: str, citations_map: Dict[str, Any]
    ) -> str:
        def replace_fn(match):
            citation_key = match.group(0)[1:-1]
            citation = citations_map.get(citation_key)
            if not citation:
                return ''
            original_display_num = citation.get('display_num')
            if original_display_num is None:
                return ''
            if original_display_num not in self._original_to_export_num:
                self._export_num_counter += 1
                self._original_to_export_num[original_display_num] = self._export_num_counter
            export_num = self._original_to_export_num[original_display_num]
            return f'[{export_num}]'

        return re.sub(r'\[citation_\d+\]', replace_fn, text)


    def _format_content_for_export(self, content: str) -> str:
        if not content:
            return content

        lines = content.split('\n')
        result_lines = []
        in_code_block = False
        in_table = False
        in_list = False

        def ensure_blank_before():
            if result_lines and result_lines[-1].strip():
                result_lines.append('')

        def is_list_line(s):
            return bool(re.match(r'^[\-\*\+]\s', s) or re.match(r'^\d+\.\s', s))

        for i, line in enumerate(lines):
            stripped = line.strip()

            if stripped.startswith('```'):
                if not in_code_block:
                    ensure_blank_before()
                in_code_block = not in_code_block
                result_lines.append(line)
                continue

            if in_code_block:
                result_lines.append(line)
                continue

            is_table_line = stripped.startswith('|') or '|---|' in stripped or '| --- |' in stripped

            if is_table_line:
                if not in_table:
                    ensure_blank_before()
                    in_table = True
                result_lines.append(line)
                continue
            elif in_table:
                in_table = False
                result_lines.append('')

            is_current_list = is_list_line(stripped)

            if is_current_list:
                if not in_list:
                    ensure_blank_before()
                    in_list = True
                result_lines.append(line)
                continue
            elif in_list:
                in_list = False
                if stripped:
                    result_lines.append('')

            if stripped.startswith('#'):
                ensure_blank_before()
                result_lines.append(line)
                continue

            if stripped.startswith('>'):
                result_lines.append(line)
                continue

            if stripped.startswith('**') and stripped.endswith('**') and len(stripped) < 50:
                ensure_blank_before()
                result_lines.append(line)
                continue

            if not stripped:
                result_lines.append(line)
                continue

            INDENT = '\u3000\u3000'
            prev_line_empty = (i == 0) or (i > 0 and not lines[i-1].strip())
            if prev_line_empty:
                result_lines.append(INDENT + stripped)
            else:
                result_lines.append(line)

        return '\n'.join(result_lines)

    def _indent_paragraph(self, text: str) -> str:
        if not text:
            return text

        INDENT = '\u3000\u3000'

        lines = text.split('\n')
        if lines:
            first_line = lines[0].lstrip()
            if first_line:
                lines[0] = INDENT + first_line

        return '\n'.join(lines)

    def _strip_italic_markers(self, text: str) -> str:
        if not text:
            return text

        pattern = r'(?<!\*)\*([^*]+)\*(?!\*)'
        return re.sub(pattern, r'\1', text)

    def _normalize_markdown_table(self, content: str) -> str:
        lines = content.strip().split('\n')
        if not lines:
            return content

        table_rows = []
        separator_found = False

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if re.match(r'^\|[\s\-:|]+\|$', line):
                separator_found = True
                table_rows.append(line)
                continue

            if '|' in line:
                if not line.startswith('|'):
                    line = '| ' + line
                if not line.endswith('|'):
                    line = line + ' |'
                table_rows.append(line)

        if not table_rows:
            return content

        max_cols = 0
        for row in table_rows:
            if not re.match(r'^\|[\s\-:|]+\|$', row):
                cells = row.split('|')[1:-1]
                max_cols = max(max_cols, len(cells))

        if max_cols == 0:
            return content

        normalized_rows = []
        for row in table_rows:
            if re.match(r'^\|[\s\-:|]+\|$', row):
                normalized_rows.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
            else:
                cells = row.split('|')[1:-1]
                while len(cells) < max_cols:
                    cells.append('')
                cells = cells[:max_cols]
                normalized_rows.append('| ' + ' | '.join(c.strip() for c in cells) + ' |')

        if not separator_found and len(normalized_rows) > 0:
            separator = '| ' + ' | '.join(['---'] * max_cols) + ' |'
            normalized_rows.insert(1, separator)

        return '\n'.join(normalized_rows)

    def _build_footnote_content(self, citation: Dict[str, Any]) -> str:
        citation_type = citation.get('type', 'segment')

        is_web_citation = (
            citation_type == 'web' or
            (citation.get('url') and not citation.get('file_name'))
        )

        if is_web_citation:
            title = citation.get('title', '未知标题')
            url = citation.get('url', '')
            source = citation.get('source', '')
            parts = [title]
            if source:
                parts.append(f"来源：{source}")
            if url:
                parts.append(f"[查看原文]({url})")
            return ' | '.join(parts)

        file_name = citation.get('file_name', '')
        page = citation.get('page') or citation.get('start_page')
        segment_id = citation.get('segment_id')

        source_parts = []
        if file_name:
            source_parts.append(f"《{file_name}》")
        if page:
            source_parts.append(f"第{page}页")

        source_str = ''.join(source_parts) or '未知来源'

        content = self._get_citation_content(citation, segment_id)

        if content:
            content = self._truncate_footnote_content(content)
            return f"{source_str}：「{content}」"

        return source_str

    def _truncate_footnote_content(self, content: str) -> str:
        if not content:
            return ''

        content = content.strip()
        if len(content) <= MAX_FOOTNOTE_CONTENT_LENGTH:
            return content

        head = content[:FOOTNOTE_HEAD_LENGTH]
        tail = content[-FOOTNOTE_TAIL_LENGTH:]
        return f"{head}……{tail}"

    def _convert_markdown_to_docx(self, markdown_content: str) -> BytesIO:
        import pypandoc

        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as md_file:
            md_file.write(markdown_content)
            md_path = md_file.name

        docx_path = md_path.replace('.md', '.docx')

        try:
            extra_args = ['--standalone']
            if self._has_reference_docx():
                extra_args.append('--reference-doc=' + self._get_reference_docx())

            pypandoc.convert_file(
                md_path,
                'docx',
                outputfile=docx_path,
                extra_args=extra_args
            )

            self._merge_duplicate_footnotes(docx_path)

            self._adjust_footnote_style(docx_path)

            with open(docx_path, 'rb') as f:
                docx_stream = BytesIO(f.read())
            docx_stream.seek(0)

            return docx_stream

        finally:
            if os.path.exists(md_path):
                os.remove(md_path)
            if os.path.exists(docx_path):
                os.remove(docx_path)

    def _merge_duplicate_footnotes(self, docx_path: str):
        import zipfile
        import shutil

        try:
            with zipfile.ZipFile(docx_path, 'r') as zf:
                if 'word/footnotes.xml' not in zf.namelist():
                    return
                document_xml = zf.read('word/document.xml').decode('utf-8')
                footnotes_xml = zf.read('word/footnotes.xml').decode('utf-8')

            fn_ref_pattern = re.compile(
                r'<w:footnoteReference\s+w:id="(\d+)"\s*/>'
            )
            all_fn_refs = list(fn_ref_pattern.finditer(document_xml))

            user_refs = []
            for m in all_fn_refs:
                word_id = m.group(1)
                if word_id in ('0', '-1'):
                    continue

                fn_ref_start = m.start()
                fn_ref_end = m.end()

                search_start = max(0, fn_ref_start - 500)
                prefix = document_xml[search_start:fn_ref_start]
                r_start_matches = list(re.finditer(r'<w:r(?:\s[^>]*)?>(?!.*<w:r(?:\s[^>]*)?>)', prefix))
                if not r_start_matches:
                    r_start_matches = list(re.finditer(r'<w:r(?:\s[^>]*)?>', prefix))

                if r_start_matches:
                    last_r_start = r_start_matches[-1]
                    r_element_start = search_start + last_r_start.start()
                else:
                    continue

                suffix = document_xml[fn_ref_end:fn_ref_end + 200]
                r_end_match = re.search(r'</w:r>', suffix)
                if r_end_match:
                    r_element_end = fn_ref_end + r_end_match.end()
                else:
                    continue

                full_r_element = document_xml[r_element_start:r_element_end]

                has_text = '<w:t>' in full_r_element or '<w:t ' in full_r_element

                user_refs.append((full_r_element, word_id, r_element_start, r_element_end, has_text))

            if not user_refs:
                return

            fn_index_to_info = {}
            for idx, (full_r_xml, word_id, start, end, has_text) in enumerate(user_refs):
                fn_index_to_info[idx] = {
                    'full_r_xml': full_r_xml,
                    'word_id': word_id,
                    'start': start,
                    'end': end,
                    'has_text': has_text
                }

            display_num_to_bookmark = {}
            bookmark_id_counter = 1000

            modified_footnotes_xml = footnotes_xml

            for fn_index, display_num in self._fn_index_to_display_num.items():
                if display_num in display_num_to_bookmark:
                    continue

                fn_info = fn_index_to_info.get(fn_index)
                if not fn_info:
                    continue

                word_id = fn_info['word_id']
                bookmark_name = f"_FnDef{display_num}"
                bookmark_id = bookmark_id_counter
                bookmark_id_counter += 1

                display_num_to_bookmark[display_num] = {
                    'name': bookmark_name,
                    'id': bookmark_id,
                    'word_id': word_id
                }

                fn_pattern = re.compile(
                    r'(<w:footnote\s+[^>]*?w:id="' + word_id + r'"[^>]*>.*?<w:p[^>]*>)',
                    re.DOTALL
                )
                match = fn_pattern.search(modified_footnotes_xml)
                if match:
                    bookmark_start = f'<w:bookmarkStart w:id="{bookmark_id}" w:name="{bookmark_name}"/>'
                    bookmark_end = f'<w:bookmarkEnd w:id="{bookmark_id}"/>'
                    insert_pos = match.end()
                    modified_footnotes_xml = (
                        modified_footnotes_xml[:insert_pos] +
                        bookmark_start + bookmark_end +
                        modified_footnotes_xml[insert_pos:]
                    )

            modified_document_xml = document_xml

            sorted_refs = sorted(fn_index_to_info.items(), key=lambda x: x[1]['start'], reverse=True)

            duplicate_fn_indices = set(item[0] for item in self._merge_list)
            duplicate_word_ids_to_delete = set()

            for fn_index, fn_info in sorted_refs:
                display_num = self._fn_index_to_display_num.get(fn_index)
                if display_num is None:
                    continue

                bookmark_info = display_num_to_bookmark.get(display_num)
                if not bookmark_info:
                    continue

                if fn_info.get('has_text'):
                    continue

                word_id = fn_info['word_id']
                bookmark_name = bookmark_info['name']

                is_duplicate = fn_index in duplicate_fn_indices

                if is_duplicate:
                    replacement_xml = (
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:fldChar w:fldCharType="begin"/>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:instrText xml:space="preserve"> NOTEREF {bookmark_name} \\h </w:instrText>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:fldChar w:fldCharType="separate"/>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:t>{display_num}</w:t>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:fldChar w:fldCharType="end"/>'
                        f'</w:r>'
                    )
                    duplicate_word_ids_to_delete.add(word_id)
                else:
                    replacement_xml = (
                        f'<w:r>'
                        f'<w:rPr><w:vanish/></w:rPr>'
                        f'<w:footnoteReference w:id="{word_id}"/>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:fldChar w:fldCharType="begin"/>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:instrText xml:space="preserve"> NOTEREF {bookmark_name} \\h </w:instrText>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:fldChar w:fldCharType="separate"/>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:t>{display_num}</w:t>'
                        f'</w:r>'
                        f'<w:r>'
                        f'<w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
                        f'<w:fldChar w:fldCharType="end"/>'
                        f'</w:r>'
                    )

                old_xml = fn_info['full_r_xml']
                if old_xml in modified_document_xml:
                    modified_document_xml = modified_document_xml.replace(old_xml, replacement_xml, 1)

            for dup_word_id in duplicate_word_ids_to_delete:
                fn_def_pattern = re.compile(
                    r'<w:footnote\s+[^>]*?w:id="' + dup_word_id + r'"[^>]*>.*?</w:footnote>',
                    re.DOTALL
                )
                modified_footnotes_xml = fn_def_pattern.sub('', modified_footnotes_xml)

            temp_dir = docx_path + '_temp'
            with zipfile.ZipFile(docx_path, 'r') as zf:
                zf.extractall(temp_dir)

            with open(os.path.join(temp_dir, 'word/footnotes.xml'), 'w', encoding='utf-8') as f:
                f.write(modified_footnotes_xml)
            with open(os.path.join(temp_dir, 'word/document.xml'), 'w', encoding='utf-8') as f:
                f.write(modified_document_xml)

            os.remove(docx_path)
            with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arc_name = os.path.relpath(file_path, temp_dir)
                        zf.write(file_path, arc_name)

            shutil.rmtree(temp_dir)

        except Exception:
            pass

    def _adjust_footnote_style(self, docx_path: str):
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(docx_path)

        try:
            if 'Normal' in doc.styles:
                normal_style = doc.styles['Normal']
                normal_style.font.size = Pt(10.5)
                normal_style.font.name = '宋体'
                normal_style.font.color.rgb = RGBColor(0, 0, 0)
                normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                normal_style.paragraph_format.line_spacing = Pt(17)
        except Exception:
            pass

        try:
            for style_name in ['List Bullet', 'List Number', 'List Paragraph']:
                if style_name in doc.styles:
                    list_style = doc.styles[style_name]
                    list_style.paragraph_format.space_before = Pt(0)
                    list_style.paragraph_format.space_after = Pt(3)
                    list_style.paragraph_format.line_spacing = Pt(17)
        except Exception:
            pass

        try:
            if 'Footnote Text' in doc.styles:
                footnote_style = doc.styles['Footnote Text']
                footnote_style.font.size = Pt(8)
                footnote_style.font.name = '宋体'
                footnote_style.font.color.rgb = RGBColor(80, 80, 80)
                footnote_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except Exception:
            pass

        try:
            if 'Hyperlink' in doc.styles:
                hyperlink_style = doc.styles['Hyperlink']
                hyperlink_style.font.color.rgb = RGBColor(0, 0, 0)
                hyperlink_style.font.underline = False
        except Exception:
            pass

        try:
            for para in doc.paragraphs:
                text = para.text.strip()
                is_list_item = (
                    text.startswith('•') or
                    text.startswith('-') or
                    text.startswith('*') or
                    (len(text) > 0 and text[0].isdigit() and ('.' in text[:3] or ')' in text[:3]))
                )

                if is_list_item:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(3)

                style_name = para.style.name if para.style else ''
                is_main_title = style_name in ('Heading 1', 'Title')
                is_source_or_disclaimer = text.startswith('来源：') or '本文档由AI生成' in text

                if is_main_title or is_source_or_disclaimer:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                para.paragraph_format.line_spacing = Pt(17)

                for run in para.runs:
                    if run.font.color.rgb is None or run.font.color.rgb == RGBColor(0, 0, 255):
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    run.italic = False
        except Exception:
            pass

        try:
            for table in doc.tables:
                self._set_table_borders(table)
                
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.font.color.rgb is None or run.font.color.rgb == RGBColor(0, 0, 255):
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                run.italic = False
        except Exception:
            pass

        doc.save(docx_path)

    def _has_reference_docx(self) -> bool:
        return os.path.exists(self._get_reference_docx())

    def _get_reference_docx(self) -> str:
        return os.path.join(os.path.dirname(__file__), '..', 'templates', 'reference.docx')

    def _is_markdown_table(self, content: str) -> bool:
        lines = content.strip().split('\n')
        if len(lines) < 2:
            return False
        for line in lines:
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                return True
        return False

    def _add_table_block(self, doc, content: str):
        lines = content.strip().split('\n')
        table_rows = []

        for line in lines:
            line = line.strip()
            if not line or not line.startswith('|'):
                continue
            if re.match(r'^\|[\s\-:|]+\|$', line):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and any(c for c in cells):
                table_rows.append(cells)

        if not table_rows:
            doc.add_paragraph(content)
            return

        if table_rows and len(table_rows[0]) > 1:
            table_rows[0] = table_rows[0][1:]

        num_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(table_rows):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = row.cells[col_idx]
                    cell.text = cell_text
                    if row_idx == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

    def _add_code_block(self, doc, content: str):
        from docx.shared import Pt, Inches

        content = content.strip()
        if content.startswith('```'):
            lines = content.split('\n')
            if len(lines) > 1:
                lines = lines[1:]
            if lines and lines[-1].strip() == '```':
                lines = lines[:-1]
            content = '\n'.join(lines)

        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(content)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

    def _get_citation_content(
        self,
        citation: Dict[str, Any],
        segment_id: Optional[str]
    ) -> str:
        content_preview = citation.get('content', '')
        if content_preview.endswith('...'):
            content_preview = content_preview[:-3]

        return self._truncate_content(content_preview)

    def _truncate_content(self, content: str) -> str:
        if not content:
            return ''

        content = content.strip()
        if len(content) > MAX_CITATION_CONTENT_LENGTH:
            return content[:MAX_CITATION_CONTENT_LENGTH] + "……（内容过长，已截断）"
        return content

    def _set_table_borders(self, table):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl = table._tbl
        tblPr = tbl.tblPr
        
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        
        borders = {
            'top': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
            'bottom': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
            'left': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
            'right': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
            'insideH': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
            'insideV': {"val": "single", "sz": "4", "space": "0", "color": "auto"}
        }
        
        for border_name, attrs in borders.items():
            existing = tblBorders.find(qn(f'w:{border_name}'))
            if existing is not None:
                tblBorders.remove(existing)
            
            border = OxmlElement(f'w:{border_name}')
            for key, value in attrs.items():
                border.set(qn(f'w:{key}'), value)
            
            tblBorders.append(border)


export_service = ExportService()
